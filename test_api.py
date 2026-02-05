import requests
from docx import Document
import io
import json

url = "http://127.0.0.1:8000/analyze/"

# Create a valid docx in memory
doc = Document()
doc.add_paragraph("John Doe")
doc.add_paragraph("Software Engineer")
doc.add_paragraph("Experience with Python, FastAPI, and React.")
doc.add_paragraph("Contact: john@example.com")

# Save to a BytesIO buffer
file_stream = io.BytesIO()
doc.save(file_stream)
file_stream.seek(0)

try:
    files = {'file': ('resume.docx', file_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {'job_description': 'Software Engineer needed with Python skills. Must have experience with FastAPI and React. This is a very long job description to pass the minimum length validation check of 50 characters required by the backend API endpoint.'}
    
    print(f"Sending request to {url}...")
    response = requests.post(url, files=files, data=data)
    print(f"Status: {response.status_code}")
    
    result = response.json()
    print(f"\n=== SUGGESTIONS ===")
    for i, s in enumerate(result.get('suggestions', []), 1):
        print(f"{i}. {s[:100]}...")
    print(f"\n=== SKILL GAP ===")
    print(f"Matched: {result.get('ats_score', {}).get('skill_gap', {}).get('matched', [])}")
    print(f"Missing: {result.get('ats_score', {}).get('skill_gap', {}).get('missing', [])}")
except Exception as e:
    print(f"Error: {e}")
